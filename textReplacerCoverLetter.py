#Create a program that takes a company, and a job title from a user, and compiles a cover letter with that information, and the correct date in both docx and pdf.
import re
from docx import Document
from docx2pdf import convert
from datetime import datetime

#Declare the filename for the Cover Letter that will be manipulated.
filename = r"C:\Users\Corey Crooks\OneDrive\zJobSearch\2026CoverLetter.docx"
document = Document(filename)

#Main replacer method.
def myreplace(file, regex, replace):
    #for all text in paragraphs,
    for p in file.paragraphs:
        #Comb each docx run, and search for the keyword to replace.
        if regex.search(p.text):
            inline=p.runs
            #For each inline word,
            for i in range (len(inline)):
                #When found a matching run,
                if regex.search(inline[i].text):
                    #Replace the keyword with the new word.
                    text=regex.sub(replace, inline[i].text)
                    inline[i].text=text
    #If the document contains a table,
    for table in file.tables:
        #Comb through each cell of the table looking for the keyword.
        for row in table.rows:
            for cell in row.cells:
                myreplace(cell,regex,replace)

#Start a helper variable to contain function calls for the date and time.
now = datetime.now()

#Set up the myreplace inputs with the placeholder {Date} found in the docx cover letter.
x = "{Date}"
#Compile into a useable string.
t = re.compile(x)
#Get the day in numbers ie 17.
day = now.strftime("%d")
#Get the month name ie April.
month = now.strftime("%B")
#Get the year in four numbers ie 2026.
year = now.strftime("%Y")
#Set the replacer variable to the full date string ie April 17, 2026.
r = month + " " + day + ", " + year

#Replace {Date} with the current date.
myreplace (document, t, r)

#Key into the {Company} placeholder.
x ="{Company}"
#Turn it into a useable string.
t = re.compile(x)
#Ask the user for the company they're applying to.
r = input('Enter the name of the Company:   ')
rNoSpace = r.replace(" ","")
#Add the company name to the end of the cover letter filepath to create a copy instead of overriding the template.
newSaveFilename = r"C:\Users\Corey Crooks\OneDrive\zJobSearch\2026CoverLetter" + rNoSpace
#Replace the {Company} placeholder with the user input.
myreplace(document, t, r)

#Key into the {Position} placeholder.
x ="{Position}"
#Compile it into a useable string.
t = re.compile(x)
#Ask the user for the position they're applying for.
r = input('Enter the position you are applying for:   ')
#Replace the {Position} placeholder with the actual position.
myreplace(document, t, r)

#Save the new docx that has been manipulated, as a new file with the company added onto the file name.
document.save(newSaveFilename + ".docx")
#Export that docx to a pdf for universal viewing.
convert(newSaveFilename+ ".docx", newSaveFilename + ".pdf")
#Notify the user that the process is complete.
print('Text is replaced, and document is saved.')